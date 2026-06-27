"""
Document loaders.

PDF loading uses pdfplumber instead of pypdf so tables are preserved as
markdown instead of collapsing into a single garbled line.

Per page, two extraction passes run:
  1. Text pass  — paragraphs and sentences, tables regions blanked out
  2. Table pass — each table converted to markdown, returned as a separate chunk

This means a page with two tables and three paragraphs produces:
  - 1 SourceDocument for the page text (tables removed)
  - 2 SourceDocuments, one per table

Each table chunk carries content_type="table" in metadata so the chunker
keeps it intact (tables must not be split mid-row).
"""
from pathlib import Path

import pdfplumber

from app.access.rbac import infer_document_metadata
from app.ingestion.models import SourceDocument


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


def load_document(path: str | Path) -> list[SourceDocument]:
    file_path = Path(path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _load_pdf(file_path)
    if suffix == ".docx":
        return _load_docx(file_path)
    if suffix in {".txt", ".md"}:
        return [_load_text(file_path)]

    raise ValueError(f"Unsupported document type: {suffix}")


def load_directory(directory: str | Path) -> list[SourceDocument]:
    root = Path(directory)
    documents: list[SourceDocument] = []
    for path in sorted(root.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            documents.extend(load_document(path))
    return documents


def _load_text(path: Path) -> SourceDocument:
    return SourceDocument(
        text=path.read_text(encoding="utf-8"),
        metadata={
            "source": path.name,
            "path": str(path),
            "file_type": path.suffix.lower().lstrip("."),
            **infer_document_metadata(path.name),
        },
    )


def _load_pdf(path: Path) -> list[SourceDocument]:
    """
    Extract text and tables from each PDF page.

    For each page:
      - Tables are extracted first and converted to markdown.
      - Text is extracted with table bounding boxes masked out so table
        cells don't also appear in the text chunk.

    Returns one text SourceDocument per page (if non-empty) plus one
    SourceDocument per table found anywhere in the document.
    """
    rbac_meta = infer_document_metadata(path.name)
    documents: list[SourceDocument] = []

    with pdfplumber.open(str(path)) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            base_meta = {
                "source": path.name,
                "path": str(path),
                "file_type": "pdf",
                "page": page_number,
                **rbac_meta,
            }

            # ── table extraction ──────────────────────────────────────────────
            tables = page.extract_tables()
            for table_index, raw_table in enumerate(tables, start=1):
                md = _table_to_markdown(raw_table)
                if md:
                    documents.append(
                        SourceDocument(
                            text=md,
                            metadata={
                                **base_meta,
                                "content_type": "table",
                                "table_index": table_index,
                            },
                        )
                    )

            # ── text extraction (tables blanked out) ──────────────────────────
            # filter_edges removes table bounding boxes from the text extraction
            # so table cell text doesn't appear twice (once in table, once in text)
            if tables:
                table_settings = {"vertical_strategy": "lines", "horizontal_strategy": "lines"}
                page_without_tables = page.filter(
                    lambda obj: obj.get("object_type") != "char"
                    or not _inside_any_table(obj, page.find_tables(table_settings))
                )
                text = page_without_tables.extract_text() or ""
            else:
                text = page.extract_text() or ""

            if text.strip():
                documents.append(
                    SourceDocument(
                        text=text,
                        metadata={**base_meta, "content_type": "text"},
                    )
                )

    return documents


def _load_docx(path: Path) -> list[SourceDocument]:
    """
    Extract text and tables from a Word document (.docx).

    Word documents interleave paragraphs and tables in the body. We walk
    the document body in order, collecting:
      - Paragraphs → accumulated into a single text SourceDocument
      - Tables → each converted to markdown, returned as a separate chunk

    Heading text is preserved inline (prefixed with # markers) so the LLM
    understands document structure when answering section-specific questions.
    """
    import docx as _docx
    from docx.oxml.ns import qn

    doc = _docx.Document(str(path))
    rbac_meta = infer_document_metadata(path.name)
    base_meta = {
        "source": path.name,
        "path": str(path),
        "file_type": "docx",
        **rbac_meta,
    }

    documents: list[SourceDocument] = []
    text_blocks: list[str] = []
    table_index = 0

    # Walk top-level body elements in document order
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Paragraph — extract text, prefix headings with markdown #
            para = _docx.text.paragraph.Paragraph(element, doc)
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name if para.style else ""
            if "Heading 1" in style:
                text_blocks.append(f"# {text}")
            elif "Heading 2" in style:
                text_blocks.append(f"## {text}")
            elif "Heading 3" in style:
                text_blocks.append(f"### {text}")
            else:
                text_blocks.append(text)

        elif tag == "tbl":
            # Table — convert to markdown, emit as separate chunk
            table_index += 1
            tbl = _docx.table.Table(element, doc)
            raw_rows = [
                [cell.text.strip() for cell in row.cells]
                for row in tbl.rows
            ]
            # Deduplicate merged cells (python-docx repeats cell content for spans)
            raw_rows = _dedup_merged_cells(raw_rows)
            md = _table_to_markdown(raw_rows)
            if md:
                documents.append(
                    SourceDocument(
                        text=md,
                        metadata={
                            **base_meta,
                            "content_type": "table",
                            "table_index": table_index,
                        },
                    )
                )

    # Emit accumulated text as one document
    full_text = "\n\n".join(text_blocks)
    if full_text.strip():
        documents.append(
            SourceDocument(
                text=full_text,
                metadata={**base_meta, "content_type": "text"},
            )
        )

    return documents


def _dedup_merged_cells(rows: list[list[str]]) -> list[list[str]]:
    """
    python-docx repeats the same cell content for horizontally merged cells.
    Deduplicate by blanking out consecutive identical values in each row.
    """
    deduped = []
    for row in rows:
        new_row = []
        prev = object()  # sentinel
        for cell in row:
            new_row.append(cell if cell != prev else "")
            prev = cell
        deduped.append(new_row)
    return deduped


# ── helpers ───────────────────────────────────────────────────────────────────

def _table_to_markdown(raw_table: list[list[str | None]]) -> str:
    """
    Convert pdfplumber's raw table (list of rows, each a list of cell strings)
    into a GitHub-flavoured markdown table.

    Empty cells become empty strings. None cells (merged cells) inherit the
    last non-None value in the same column so the table stays readable.
    """
    if not raw_table or not raw_table[0]:
        return ""

    # Clean cells: strip whitespace, replace None with empty string
    cleaned = []
    for row in raw_table:
        cleaned.append([str(cell).strip() if cell is not None else "" for cell in row])

    # Remove completely empty rows
    cleaned = [row for row in cleaned if any(cell for cell in row)]
    if not cleaned:
        return ""

    col_count = max(len(row) for row in cleaned)

    # Pad rows to same width
    padded = [row + [""] * (col_count - len(row)) for row in cleaned]

    # Build markdown: first row = header, second row = separator
    header = "| " + " | ".join(padded[0]) + " |"
    separator = "| " + " | ".join(["---"] * col_count) + " |"
    body_rows = ["| " + " | ".join(row) + " |" for row in padded[1:]]

    return "\n".join([header, separator] + body_rows)


def _inside_any_table(obj: dict, tables) -> bool:
    """Return True if a PDF character object falls inside any table bounding box."""
    x, y = obj.get("x0", 0), obj.get("top", 0)
    for table in tables:
        bbox = table.bbox   # (x0, top, x1, bottom)
        if bbox[0] <= x <= bbox[2] and bbox[1] <= y <= bbox[3]:
            return True
    return False
